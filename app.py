from flask import Flask, render_template, send_file, jsonify, request
from google.oauth2 import service_account
from googleapiclient.discovery import build
from collections import OrderedDict
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
import json
import tempfile
from datetime import datetime
import traceback
import logging

# 로깅 설정
logging.basicConfig(level=logging.INFO)

app = Flask(__name__)

SCOPES = ['https://www.googleapis.com/auth/spreadsheets']
SPREADSHEET_ID = '1eAM2Yt9PxbY_rtKaBqf0SazNCV6j7NCw21Ho-yIcVOs'
SHEET_NAME = 'published'

CATEGORY_ORDER = ['도서', '수상', '특허', 'SW등록', '국제논문', '국내논문', '국제컨퍼런스', '국내컨퍼런스']

def get_google_sheets_service():
    try:
        app.logger.info("Attempting to create Google Sheets service...")
        service_account_info = json.loads(os.environ.get('GOOGLE_CREDENTIALS', '{}'))
        
        if not service_account_info:
            app.logger.error("No Google service account credentials found in environment variables")
            raise ValueError("No Google service account credentials found in environment variables")
            
        creds = service_account.Credentials.from_service_account_info(
            service_account_info,
            scopes=SCOPES
        )
        service = build('sheets', 'v4', credentials=creds)
        app.logger.info("Google Sheets service created successfully")
        return service
    except Exception as e:
        app.logger.error(f"Error creating Google Sheets service: {str(e)}")
        traceback.print_exc()
        return None

def get_sheet_data():
    try:
        app.logger.info("Fetching sheet data...")
        service = get_google_sheets_service()
        if not service:
            return []
        
        result = service.spreadsheets().values().get(
            spreadsheetId=SPREADSHEET_ID,
            range=f'{SHEET_NAME}!A:J'
        ).execute()
        
        data = result.get('values', [])
        app.logger.info(f"Successfully fetched {len(data)} rows of data")
        return data
    except Exception as e:
        app.logger.error(f"Error fetching sheet data: {str(e)}")
        traceback.print_exc()
        return []

def process_data(data):
    try:
        app.logger.info("Processing data...")
        organized_data = OrderedDict()
        for category in CATEGORY_ORDER:
            organized_data[category] = []
        
        rows = data[1:] if len(data) > 1 else []
        rows.sort(key=lambda x: x[3] if len(x) > 3 else '0', reverse=True)
        
        for row in rows:
            if len(row) < 2:
                continue
                
            category = row[1].strip()
            if category not in CATEGORY_ORDER:
                continue

            # Split acknowledgement by commas and trim each part
            acknowledgement = row[9] if len(row) > 9 else ''
            if acknowledgement:
                acknowledgement_parts = [part.strip() for part in acknowledgement.split(',')]
                acknowledgement = ', '.join(acknowledgement_parts)

            entry = {
                'year': row[3] if len(row) > 3 else '',
                'author': row[2] if len(row) > 2 else '',
                'title': row[4] if len(row) > 4 else '',
                'journal': row[5] if len(row) > 5 else '',
                'volume': row[6] if len(row) > 6 else '',
                'pages': row[7] if len(row) > 7 else '',
                'doi': row[8] if len(row) > 8 else '',
                'acknowledgement': acknowledgement
            }
            
            organized_data[category].append(entry)
        
        for category in organized_data:
            organized_data[category].sort(key=lambda x: x['year'] if x['year'] else '0', reverse=True)
            for idx, entry in enumerate(organized_data[category], 1):
                entry['number'] = str(idx)

        app.logger.info("Data processing completed successfully")
        return organized_data
    except Exception as e:
        app.logger.error(f"Error processing data: {str(e)}")
        traceback.print_exc()
        return {}

def apply_filters(entry, category, filters):
    if not filters:
        return True
        
    if filters.get('text'):
        searchable_text = f"{category} {entry['author']} {entry['title']} {entry['journal']} {entry['volume']} {entry['pages']} {entry['doi']} {entry['acknowledgement']}".lower()
        if filters['text'].lower() not in searchable_text:
            return False
    
    if filters.get('categories'):
        if category not in filters['categories']:
            return False
    
    if filters.get('years'):
        if entry['year'] not in filters['years']:
            return False
    
    if filters.get('fundings'):
        acknowledgement_parts = [part.strip() for part in entry['acknowledgement'].split(',')]
        if not any(funding in acknowledgement_parts for funding in filters['fundings']):
            return False
    
    return True

def filter_data(organized_data, filters):
    try:
        app.logger.info("Applying filters to data...")
        filtered_data = OrderedDict()
        
        for category, entries in organized_data.items():
            filtered_entries = []
            for entry in entries:
                if apply_filters(entry, category, filters):
                    filtered_entries.append(entry)
            if filtered_entries:
                filtered_data[category] = filtered_entries
                
        app.logger.info("Filters applied successfully")
        return filtered_data
    except Exception as e:
        app.logger.error(f"Error filtering data: {str(e)}")
        traceback.print_exc()
        return {}

def create_word_document(data):
    try:
        app.logger.info("Creating Word document...")
        doc = Document()
        for section in doc.sections:
            section.page_margin_top = Inches(1)
            section.page_margin_bottom = Inches(1)
            section.page_margin_left = Inches(1)
            section.page_margin_right = Inches(1)
        
        # 임시 파일 생성
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            filepath = tmp.name
            
            for category in data.keys():
                if data[category]:
                    doc.add_heading(category, level=1)
                    
                    for entry in data[category]:
                        p = doc.add_paragraph()
                        p.style = 'List Bullet'
                        
                        if category == '도서':
                            text = f"{entry['author']}. ({entry['year']}) \"{entry['title']}\", {entry['journal']}, {entry['volume']}"
                            if entry['doi']:
                                clean_doi = entry['doi'].strip().replace('https://doi.org/', '')
                                text += f". https://doi.org/{clean_doi}"
                            p.add_run(text)
                        
                        elif category == '수상':
                            text = f"{entry['journal']}, {entry['author']}, {entry['title']}, {entry['year']}"
                            p.add_run(text)
                        
                        elif category == '특허':
                            text = f"{entry['title']}, 출원번호: {entry['journal']}, {entry['year']}"
                            p.add_run(text)
                        
                        elif category == 'SW등록':
                            text = f"{entry['title']}, 등록번호: {entry['journal']}, {entry['year']}"
                            p.add_run(text)
                        
                        else:  
                            p.add_run(f"{entry['author']}. ({entry['year']}). {entry['title']}. ")
                            
                            journal_run = p.add_run(entry['journal'])
                            journal_run.italic = True
                            
                            if entry['volume'] or entry['pages']:
                                if entry['volume']:
                                    vol_parts = entry['volume'].split('(')
                                    if len(vol_parts) > 1:
                                        p.add_run(", ")
                                        vol_run = p.add_run(vol_parts[0])
                                        vol_run.italic = True
                                        p.add_run(f"({vol_parts[1]}")
                                    else:
                                        p.add_run(", ")
                                        vol_run = p.add_run(entry['volume'])
                                        vol_run.italic = True
                                if entry['pages']:
                                    p.add_run(f", {entry['pages']}")
                            
                            if entry['doi']:
                                clean_doi = entry['doi'].strip().replace('https://doi.org/', '')
                                p.add_run(f". https://doi.org/{clean_doi}")
                    
                    doc.add_paragraph()

            doc.save(filepath)
            app.logger.info("Word document created successfully")
            return filepath
    except Exception as e:
        app.logger.error(f"Error creating Word document: {str(e)}")
        traceback.print_exc()
        return None

@app.route('/export/word', methods=['POST'])
def export_word():
    try:
        app.logger.info("Processing word export request...")
        data = request.json
        filename = create_word_document(data)
        if not filename:
            app.logger.error("Failed to create Word document")
            return "Word 문서 생성에 실패했습니다.", 500

        response = send_file(
            filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            download_name='DesignITLab_CV.docx',
            as_attachment=True
        )
        app.logger.info("Word document exported successfully")
        return response
    except Exception as e:
        app.logger.error(f"Error exporting Word: {str(e)}")
        traceback.print_exc()
        return "Word 문서 내보내기에 실패했습니다.", 500
    finally:
        if 'filename' in locals() and os.path.exists(filename):
            try:
                os.remove(filename)
                app.logger.info("Temporary file removed successfully")
            except Exception as e:
                app.logger.error(f"Error removing temporary file: {str(e)}")

@app.route('/')
def index():
    try:
        app.logger.info("Processing index route request...")
        sheet_data = get_sheet_data()
        organized_data = process_data(sheet_data)
        
        filters = {
            'text': request.args.get('search', ''),
            'categories': request.args.getlist('category'),
            'years': request.args.getlist('year'),
            'fundings': request.args.getlist('funding')
        }
        
        if any(filters.values()):
            app.logger.info("Applying filters to data...")
            filtered_data = filter_data(organized_data, filters)
            return render_template('index.html', categories=filtered_data)
        
        app.logger.info("Rendering index template with all data")
        return render_template('index.html', categories=organized_data)
    except Exception as e:
        app.logger.error(f"Error in index route: {str(e)}")
        traceback.print_exc()
        return render_template('index.html', categories={})

# 에러 핸들러 추가
@app.errorhandler(404)
def not_found_error(error):
    app.logger.error(f"Page not found: {request.url}")
    return render_template('404.html'), 404

@app.errorhandler(500)
def internal_error(error):
    app.logger.error(f"Server Error: {error}")
    return render_template('500.html'), 500

# Render 환경을 위한 포트 설정
port = int(os.environ.get('PORT', 10000))

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=port)
else:
    # production 환경에서는 gunicorn이 이 부분을 처리
    gunicorn_logger = logging.getLogger('gunicorn.error')
    app.logger.handlers = gunicorn_logger.handlers
    app.logger.setLevel(gunicorn_logger.level)
